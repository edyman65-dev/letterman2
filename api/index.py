import os, re, io, json, zipfile
from datetime import datetime
from flask import Flask, request, jsonify, send_file
import pandas as pd
from docx import Document

app = Flask(__name__)
app.config['MAX_CONTENT_LENGTH'] = 50 * 1024 * 1024

W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

# ── Supabase (base de dados gratuita) ────────────────────────
from supabase import create_client

SUPABASE_URL = os.environ.get('SUPABASE_URL', '')
SUPABASE_KEY = os.environ.get('SUPABASE_KEY', '')

def get_db():
    return create_client(SUPABASE_URL, SUPABASE_KEY)

def load_enderecos():
    try:
        db   = get_db()
        rows = db.table('enderecos').select('*').execute().data
        return {r['empresa']: {'avenida': r['avenida'], 'numero': r['numero'], 'postal': r['postal']} for r in rows}
    except Exception:
        return {}

def save_enderecos(data: dict):
    try:
        db = get_db()
        for empresa, end in data.items():
            db.table('enderecos').upsert({
                'empresa': empresa,
                'avenida': end.get('avenida', ''),
                'numero':  end.get('numero',  ''),
                'postal':  end.get('postal',  ''),
            }).execute()
    except Exception as e:
        print(f'Erro ao guardar endereços: {e}')

# ── Seed dos endereços conhecidos ────────────────────────────
ENDERECOS_SEED = {
    "Africa Communications Lda":                        {"avenida": "Av 24 de julho",            "numero": "2096", "postal": ""},
    "Altel Soluções Globais":                           {"avenida": "AV.FPLM",                   "numero": "",     "postal": "1108"},
    "Amani Shared Services Lda":                        {"avenida": "Av Kim Il Sung",             "numero": "65",   "postal": ""},
    "Auto Sueco Moçambique S.A":                        {"avenida": "Av da Namaacha",             "numero": "8274", "postal": ""},
    "CLIDIS Lda":                                       {"avenida": "Av Dr Nkutumula",            "numero": "511",  "postal": ""},
    "Cargofrete, Lda":                                  {"avenida": "Av- do Aeroporto",           "numero": "",     "postal": "3200"},
    "Comserv Moçambique.L":                             {"avenida": "Rua Crisanto Castiano",      "numero": "",     "postal": "1100"},
    "Dalima, Lda":                                      {"avenida": "Av 25 de Setembro",          "numero": "",     "postal": "1100"},
    "EDIMADE CONSTRUÇÃO & IMOBILIÁRIO MOÇAMBI":         {"avenida": "Rua da Resistência",         "numero": "",     "postal": "1106"},
    "EDIVISA - Empresa de Construções SA":              {"avenida": "Av do Trabalho",             "numero": "1501", "postal": ""},
    "Embaixada da Argelia em Maputo":                   {"avenida": "Rua de Mukumbura",           "numero": "",     "postal": ""},
    "IDE":                                              {"avenida": "Av do Zimbabwe",             "numero": "1688", "postal": ""},
    "IMOVISA - IMOBILIÁRIA DE MOÇAMBIQUE S.A":          {"avenida": "Av 25 de Setembro",          "numero": "111",  "postal": ""},
    "Laboratório Joaquim Chaves Moçambique,":           {"avenida": "Av 24 de Julho",             "numero": "",     "postal": "1100"},
    "MSC Moçambique, LDA":                              {"avenida": "Rua dos desportistas",       "numero": "833",  "postal": ""},
    "NORS Moçambique S.A.":                             {"avenida": "Av da Namaacha",             "numero": "8274", "postal": ""},
    "PETROGAL MOÇAMBIQUE, LDA":                         {"avenida": "Rua dos desportistas",       "numero": "83",   "postal": "1100"},
    "Painhas, SA":                                      {"avenida": "Rua dos desportistas",       "numero": "",     "postal": ""},
    "Portucel Moçambique SA.":                          {"avenida": "Av Marginal",                "numero": "141",  "postal": ""},
    "SGS MCnet":                                        {"avenida": "Av da Namaacha",             "numero": "8274", "postal": ""},
    "SGS Moçambique, Lda":                              {"avenida": "Av da Namaacha",             "numero": "8274", "postal": ""},
    "Salamanga Comercial, SA":                          {"avenida": "Emília Daússe",              "numero": "",     "postal": "1100"},
    "Star Gás":                                         {"avenida": "Av Moçambique",              "numero": "4830", "postal": ""},
    "Tecnel Service, Lda":                              {"avenida": "Av das Industrias",          "numero": "760",  "postal": ""},
    "Televisa - Sociedade Técnica de Obras e Projectos":{"avenida": "Av dos Presidentes",         "numero": "",     "postal": ""},
    "ZACI,SA":                                          {"avenida": "Av Kim Il Sung",             "numero": "",     "postal": "1102"},
    "Marin Trading, Lda":                               {"avenida": "Av Angola",                  "numero": "",     "postal": "1100"},
    "MOZA BANCO, SA":                                   {"avenida": "Rua dos desportistas",       "numero": "713",  "postal": "1100"},
    "Fundo da Economia Azul - ProAzul, FP":             {"avenida": "Emília Daússe",              "numero": "",     "postal": "1100"},
    "MANUPLAN - Sociedade Unipessoal, Lda":             {"avenida": "Rua Marconi",                "numero": "84",   "postal": ""},
    "Executive Protetion":                              {"avenida": "Av Mao Tsé Tung",            "numero": "1483", "postal": ""},
    "PIONEIRA ALIMENTAR MOÇAMBICANA":                   {"avenida": "Av do Trabalho",             "numero": "1743", "postal": ""},
    "Universidade Save - Extensão de Maxixe":           {"avenida": "Estrada Nacional",           "numero": "12",   "postal": ""},
    "Transauto, Lda":                                   {"avenida": "Av Josina Machel",           "numero": "1149", "postal": ""},
    "Mota-Engil Engenharia":                            {"avenida": "Rua do Kassuende",           "numero": "210",  "postal": ""},
    "ENGTEL MOÇAMBIQUE SA":                             {"avenida": "Rua do Comércio",            "numero": "132",  "postal": ""},
    "TDGI Lda":                                         {"avenida": "Av 24 de Julho",             "numero": "11",   "postal": "1100"},
    "Caetano Equipamentos SA":                          {"avenida": "Av Angola",                  "numero": "2290", "postal": ""},
    "Fidelidade Moçambique, SA":                        {"avenida": "Rua do Kassuende",           "numero": "",     "postal": ""},
    "General betting, LDA":                             {"avenida": "Av. Marien Ngoabi",          "numero": "49",   "postal": ""},
    "Entreposto Comercial de Moç. S.A.":                {"avenida": "Av do trabalho",             "numero": "1856", "postal": "1100"},
    "Mediplus Companhia de Seguros, SA":                {"avenida": "Bairro da Coop",             "numero": "27",   "postal": "1100"},
    "Moçambique Terramar Trading, Lda":                 {"avenida": "Av de Moçambique",           "numero": "1211", "postal": ""},
    "Yunike Mkt, SA":                                   {"avenida": "Av Mártires da Machava",     "numero": "",     "postal": "1102"},
    "Bravantic Envolving":                              {"avenida": "Av 24 de Julho",             "numero": "11",   "postal": ""},
    "Arkhe Risk Solutions Lda":                         {"avenida": "Daniel Napatima",            "numero": "",     "postal": "1102"},
    "IMOPETRO-Importadora Moç. de Petróleo":            {"avenida": "25 de Setembro",             "numero": "1230", "postal": ""},
    "CFM Logistics S,A":                                {"avenida": "Rua dos desportistas",       "numero": "480",  "postal": ""},
    "Lancet Laboratories Mozambique, Lda":              {"avenida": "Av Palmar",                  "numero": "238",  "postal": ""},
    "Tintas CIN de Moçambique, SA":                     {"avenida": "Av 24 de Julho",             "numero": "3736", "postal": ""},
    "Moçambique Florestal, S.A":                        {"avenida": "Rua Antonio Enes",           "numero": "76",   "postal": ""},
    "TVCabo-Comunicações Multimédia, Lda":              {"avenida": "Av dos Presidentes",         "numero": "1100", "postal": ""},
    "Sociedade Industrial de Pesca, Lda":               {"avenida": "Porto de pesca de Maputo",   "numero": "",     "postal": ""},
    "Electrotec S.A.":                                  {"avenida": "Av de Moçambique",           "numero": "38",   "postal": ""},
    "Técnica Industrial, SARL":                         {"avenida": "Av de Angola",               "numero": "2119", "postal": ""},
}

# ── Lógica de geração ────────────────────────────────────────

def detect_tpl_info(doc):
    info = {"nome_placeholder": None, "nr_unico_no": False}
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    nos   = [t.text for t in p._element.iter(f'{{{W}}}t') if t.text]
                    texto = ''.join(nos)
                    if 'nome' in texto:
                        for t in p._element.iter(f'{{{W}}}t'):
                            if t.text and 'nome' in t.text:
                                info["nome_placeholder"] = t.text
                    if 'Nr.' in texto and 'numero' in texto:
                        info["nr_unico_no"] = len(nos) == 1
    return info

def sub_xml(elem, refs):
    for t in elem.iter(f'{{{W}}}t'):
        if not t.text: continue
        for c, v in refs.items():
            if c in t.text:
                t.text = t.text.replace(c, str(v))

def gerar_carta(tpl_bytes, tpl_info, nome, mat, av, nr, post, serie, ref, DD, MM, YY):
    doc      = Document(io.BytesIO(tpl_bytes))
    nome_key = tpl_info["nome_placeholder"]
    nome_val = (nome_key[0] if nome_key and nome_key[0] == ' ' else '') + nome
    refs = {nome_key: nome_val, 'matricula': mat, 'avenida': av, 'postal': post,
            'serie': serie, 'k': ref, 'DD': DD, 'MM': MM, 'YY': YY}
    for p in doc.paragraphs:
        sub_xml(p._element, refs)
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    nos_texto = ''.join(t.text for t in p._element.iter(f'{{{W}}}t') if t.text)
                    if 'Nr.' in nos_texto and 'numero' in nos_texto:
                        if tpl_info["nr_unico_no"]:
                            for t in p._element.iter(f'{{{W}}}t'):
                                if t.text and 'Nr.' in t.text and 'numero' in t.text:
                                    t.text = t.text.replace('numero', nr.strip()) if nr.strip() else ''
                        else:
                            for t in p._element.iter(f'{{{W}}}t'):
                                txt = t.text or ''
                                if 'Nr. ' in txt:
                                    t.text = '' if not nr.strip() else txt
                                elif 'numero' in txt:
                                    t.text = txt.replace('numero', nr.strip())
                    else:
                        sub_xml(p._element, refs)
    for shape in doc.element.iter():
        if shape.tag.endswith('txbxContent'):
            sub_xml(shape, refs)
    out = io.BytesIO()
    doc.save(out)
    out.seek(0)
    return out.read()

def extrair(nr):
    nr = str(nr).strip().replace(' ', '')
    if len(nr) < 18: return nr, ''
    return nr[:6] + nr[-6:], nr[6:12]

def tipo_carta(t, r):
    t = str(t).upper() if pd.notna(t) else ''
    r = str(r).upper() if pd.notna(r) else ''
    tag = 'TAG' in t
    rev = 'REVALIDA' in r or 'REVALIDA' in t
    if tag and rev: return 'revalidacao_com_tag'
    if tag:         return 'cartao_com_tag'
    if rev:         return 'revalidacao_sem_tag'
    return 'cartao_sem_tag'

# ── Rotas ─────────────────────────────────────────────────────

@app.route('/api/seed', methods=['POST'])
def seed():
    existing = load_enderecos()
    to_add   = {k: v for k, v in ENDERECOS_SEED.items() if k not in existing}
    if to_add:
        save_enderecos(to_add)
    return jsonify({'ok': True, 'seed': len(to_add), 'total': len(existing) + len(to_add)})

@app.route('/api/enderecos', methods=['GET'])
def get_enderecos():
    return jsonify(load_enderecos())

@app.route('/api/enderecos', methods=['POST'])
def update_enderecos():
    data = request.json or {}
    save_enderecos(data)
    return jsonify({'ok': True})

@app.route('/api/analisar', methods=['POST'])
def analisar():
    if 'principal' not in request.files:
        return jsonify({'erro': 'Ficheiro principal em falta'}), 400
    lote = int(request.form.get('lote', 9))
    df   = pd.read_excel(request.files['principal'], sheet_name='Dados', dtype=str)
    df.columns = df.columns.str.strip()
    col_obs = next((c for c in df.columns if 'bserva' in c), None)
    if not col_obs:
        return jsonify({'erro': 'Coluna Observação não encontrada'}), 400
    filtro  = df[col_obs].str.strip().str.lower() == f'lote {lote}'.lower()
    df_lote = df[filtro].copy()
    if len(df_lote) == 0:
        disp = sorted(df[df[col_obs].str.contains('Lote', na=False, case=False)][col_obs].str.strip().unique().tolist())
        return jsonify({'erro': f'Lote {lote} não encontrado', 'disponiveis': disp}), 404
    col_tipo    = next((c for c in df.columns if 'série' in c.lower() or 'serie' in c.lower()), None)
    col_revalid = 'Letra gravada'
    contagem    = {'cartao_com_tag': 0, 'cartao_sem_tag': 0, 'revalidacao_com_tag': 0, 'revalidacao_sem_tag': 0}
    for _, row in df_lote.iterrows():
        contagem[tipo_carta(row.get(col_tipo, ''), row.get(col_revalid, ''))] += 1
    enderecos    = load_enderecos()
    empresas     = df_lote['Cliente'].dropna().unique().tolist()
    sem_endereco = [e for e in empresas if e not in enderecos]
    return jsonify({'total': len(df_lote), 'empresas': empresas,
                    'sem_endereco': sem_endereco, 'contagem': contagem, 'lote': lote})

@app.route('/api/gerar', methods=['POST'])
def gerar():
    for k in ['principal', 'tpl_com_tag', 'tpl_sem_tag', 'tpl_reval_com', 'tpl_reval_sem']:
        if k not in request.files:
            return jsonify({'erro': f'{k} em falta'}), 400
    lote  = int(request.form.get('lote', 9))
    novos = json.loads(request.form.get('enderecos_novos', '{}'))
    if novos:
        save_enderecos(novos)
    enderecos = load_enderecos()
    df        = pd.read_excel(request.files['principal'], sheet_name='Dados', dtype=str)
    df.columns = df.columns.str.strip()
    col_obs      = next((c for c in df.columns if 'bserva' in c), None)
    col_cartao   = next((c for c in df.columns if 'cart' in c.lower()), None)
    col_matricula= next((c for c in df.columns if 'matr' in c.lower() or 'nome/m' in c.lower()), None)
    col_tipo     = next((c for c in df.columns if 'série' in c.lower() or 'serie' in c.lower()), None)
    col_revalid  = 'Letra gravada'
    filtro       = df[col_obs].str.strip().str.lower() == f'lote {lote}'.lower()
    df_lote      = df[filtro].copy()
    tpl_map = {
        'cartao_com_tag':      request.files['tpl_com_tag'].read(),
        'cartao_sem_tag':      request.files['tpl_sem_tag'].read(),
        'revalidacao_com_tag': request.files['tpl_reval_com'].read(),
        'revalidacao_sem_tag': request.files['tpl_reval_sem'].read(),
    }
    tpl_info = {k: detect_tpl_info(Document(io.BytesIO(v))) for k, v in tpl_map.items()}
    hoje = datetime.now()
    DD, MM, YY = str(hoje.day).zfill(2), str(hoje.month).zfill(2), str(hoje.year)[-2:]
    zip_buf = io.BytesIO()
    cnt     = {k: 0 for k in tpl_map}
    with zipfile.ZipFile(zip_buf, 'w', zipfile.ZIP_DEFLATED) as zf:
        for _, row in df_lote.iterrows():
            nr_c  = str(row[col_cartao]).strip()
            cli   = str(row['Cliente']).strip()
            mat   = str(row[col_matricula]).strip() if pd.notna(row[col_matricula]) else ''
            tipo  = tipo_carta(row.get(col_tipo, ''), row.get(col_revalid, ''))
            serie, ref = extrair(nr_c)
            end   = enderecos.get(cli, {})
            cnt[tipo] += 1
            ns    = re.sub(r'[\\/*?:"<>|]', '', cli[:25]).strip()
            carta = gerar_carta(tpl_map[tipo], tpl_info[tipo], cli, mat,
                                end.get('avenida',''), end.get('numero',''),
                                end.get('postal',''), serie, ref, DD, MM, YY)
            zf.writestr(f'L{lote}_{tipo}_{cnt[tipo]:03d}_{ns}.docx', carta)
    zip_buf.seek(0)
    return send_file(zip_buf, mimetype='application/zip', as_attachment=True,
                     download_name=f'Cartas_Lote{lote}.zip')
